"""Word (.docx) report generation for per-fund factsheet and comparison.

Two entry points:
- build_fund_factsheet(session, scheme_code) -> bytes
- build_comparison_report(session, scheme_codes) -> bytes

Branding: Z1N Capital, Arial 11pt body, teal accent (#0F766E).
Page: A4 portrait, 1" margins.

Charts: embedded as PNG via chart_render. Server-side, no frontend dep.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import desc, select
from sqlalchemy.orm import Session

from app.models.fund import Fund, FundMetric, FundScore, NavHistory
from app.services.chart_render import render_nav_chart, render_overlay_chart

logger = logging.getLogger(__name__)

BRAND_TEAL = RGBColor(0x0F, 0x76, 0x6E)
GREY = RGBColor(0x64, 0x74, 0x8B)
DARK = RGBColor(0x0F, 0x17, 0x2A)
LIGHT_BG = "F0FDF4"  # cell shade for header rows


# ---- low-level helpers ---------------------------------------------------

def _set_cell_shade(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_heading(doc, text: str, size: int = 14, color: RGBColor = BRAND_TEAL) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color


def _add_para(doc, text: str, size: int = 10, color: RGBColor = DARK, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _make_kv_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(6.5)
    table.columns[1].width = Cm(8.5)
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(6.5)
        c1.width = Cm(8.5)
        _set_cell_shade(c0, LIGHT_BG)
        c0.paragraphs[0].add_run(k).font.bold = True
        c0.paragraphs[0].runs[0].font.name = "Arial"
        c0.paragraphs[0].runs[0].font.size = Pt(9)
        c1.paragraphs[0].add_run(v).font.name = "Arial"
        c1.paragraphs[0].runs[0].font.size = Pt(9)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def _make_grid_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = False
    # Header row
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        _set_cell_shade(c, "0F766E")
        run = c.paragraphs[0].add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            run = c.paragraphs[0].add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(9)
    return table


def _add_image_bytes(doc, png: bytes, width_cm: float = 16.0) -> None:
    doc.add_picture(io.BytesIO(png), width=Cm(width_cm))


def _add_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Z1N Capital - Internal research. Not investment advice. Past performance is not "
        "indicative of future returns."
    )
    run.font.name = "Arial"
    run.font.size = Pt(7)
    run.font.color.rgb = GREY


def _format_pct(v, digits: int = 2) -> str:
    if v is None:
        return "-"
    return f"{v * 100:.{digits}f}%"


def _format_num(v, digits: int = 2) -> str:
    if v is None:
        return "-"
    return f"{v:.{digits}f}"


# ---- data loaders --------------------------------------------------------

def _load_fund_bundle(session: Session, scheme_code: int) -> dict | None:
    fund = session.get(Fund, scheme_code)
    if fund is None:
        return None
    metric = session.get(FundMetric, scheme_code)
    score = session.execute(
        select(FundScore)
        .where(FundScore.scheme_code == scheme_code)
        .order_by(desc(FundScore.computed_at))
        .limit(1)
    ).scalar_one_or_none()
    nav_rows = session.execute(
        select(NavHistory.nav_date, NavHistory.nav)
        .where(NavHistory.scheme_code == scheme_code)
        .order_by(NavHistory.nav_date)
    ).all()
    latest = nav_rows[-1] if nav_rows else None
    return {
        "fund": fund,
        "metric": metric,
        "score": score,
        "nav_history": [(d, float(n)) for (d, n) in nav_rows],
        "latest_nav": float(latest.nav) if latest else None,
        "nav_date": latest.nav_date if latest else None,
    }


# ---- factsheet -----------------------------------------------------------

def build_fund_factsheet(session: Session, scheme_code: int) -> bytes:
    """Build a per-fund factsheet .docx. Returns raw bytes."""
    bundle = _load_fund_bundle(session, scheme_code)
    if bundle is None:
        raise ValueError(f"Fund {scheme_code} not found")

    fund = bundle["fund"]
    metric = bundle["metric"]
    score = bundle["score"]

    doc = Document()
    # Page margins (A4 default in python-docx).
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Header strip
    head = doc.add_paragraph()
    r = head.add_run("Z1N CAPITAL")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = BRAND_TEAL
    _add_heading(doc, fund.fund_name, size=18)
    _add_para(doc, f"{fund.amc or '-'} | {fund.category or '-'} | {fund.sub_category or '-'}",
              size=10, color=GREY)

    # Overview key-value
    _add_heading(doc, "Snapshot", size=12)
    _make_kv_table(doc, [
        ("Latest NAV", f"Rs {bundle['latest_nav']:.4f}" if bundle["latest_nav"] else "-"),
        ("As of", bundle["nav_date"].isoformat() if bundle["nav_date"] else "-"),
        ("Composite score (0-100)", _format_num(score.composite_score if score else None, 1)),
        ("Expense ratio (annual)", f"{fund.expense_ratio:.2f}%" if fund.expense_ratio else "-"),
        ("Exit load", fund.exit_load or "None"),
        ("AUM (Rs crore)", f"{fund.aum_cr:,.0f}" if fund.aum_cr else "-"),
        ("Plan type", fund.plan_type or "-"),
    ])

    # Returns
    _add_heading(doc, "Annualised returns", size=12)
    _make_grid_table(
        doc,
        ["1 year", "3 year", "5 year", "10 year"],
        [[
            _format_pct(metric.cagr_1y if metric else None),
            _format_pct(metric.cagr_3y if metric else None),
            _format_pct(metric.cagr_5y if metric else None),
            _format_pct(metric.cagr_10y if metric else None),
        ]],
    )

    # Risk panel
    _add_heading(doc, "Risk profile", size=12)
    _make_grid_table(
        doc,
        ["Sharpe ratio", "Volatility (std dev)", "Worst drawdown", "Recovery (months)"],
        [[
            _format_num(metric.sharpe_ratio if metric else None, 2),
            _format_pct(metric.std_dev if metric else None, 1),
            _format_pct(metric.max_drawdown if metric else None, 1),
            str(metric.recovery_months) if metric and metric.recovery_months else "-",
        ]],
    )

    # Score breakdown
    if score:
        _add_heading(doc, "Score breakdown", size=12)
        _make_grid_table(
            doc,
            ["Sharpe", "CAGR 1y", "CAGR 3y", "CAGR 5y", "Drawdown", "Expense", "Momentum"],
            [[
                _format_num(score.sharpe_score, 0),
                _format_num(score.cagr_1y_score, 0),
                _format_num(score.cagr_3y_score, 0),
                _format_num(score.cagr_5y_score, 0),
                _format_num(score.drawdown_score, 0),
                _format_num(score.expense_score, 0),
                _format_num(score.momentum_score, 0),
            ]],
        )

    # NAV chart
    _add_heading(doc, "NAV history", size=12)
    nav_png = render_nav_chart(
        bundle["nav_history"],
        title="NAV over time",
        cache_key=f"fund_{scheme_code}_{bundle['nav_date'] or 'na'}",
    )
    _add_image_bytes(doc, nav_png, width_cm=16.0)

    # Footer with disclaimer + generation timestamp
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _add_para(doc, f"Generated {ts}", size=8, color=GREY)
    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- comparison report ---------------------------------------------------

def build_comparison_report(session: Session, scheme_codes: list[int]) -> bytes:
    """Build a 2-5 fund side-by-side comparison .docx. Returns raw bytes."""
    if not 1 <= len(scheme_codes) <= 5:
        raise ValueError("Provide 1-5 scheme_codes for comparison")
    bundles = [_load_fund_bundle(session, c) for c in scheme_codes]
    if any(b is None for b in bundles):
        missing = [c for c, b in zip(scheme_codes, bundles, strict=False) if b is None]
        raise ValueError(f"Schemes not found: {missing}")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    head = doc.add_paragraph()
    r = head.add_run("Z1N CAPITAL")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = BRAND_TEAL
    _add_heading(doc, "Fund Comparison Report", size=18)
    names_line = ", ".join(b["fund"].fund_name for b in bundles)
    _add_para(doc, names_line, size=9, color=GREY)

    # Side-by-side matrix
    headers = ["Metric"] + [b["fund"].fund_name[:30] for b in bundles]

    def _row(label, vals):
        return [label, *vals]

    rows = [
        _row("AMC", [b["fund"].amc or "-" for b in bundles]),
        _row("Category", [b["fund"].category or "-" for b in bundles]),
        _row("Sub-category", [b["fund"].sub_category or "-" for b in bundles]),
        _row("Composite score", [
            _format_num(b["score"].composite_score, 1) if b["score"] else "-" for b in bundles
        ]),
        _row("Expense ratio", [
            f"{b['fund'].expense_ratio:.2f}%" if b["fund"].expense_ratio else "-" for b in bundles
        ]),
        _row("AUM (Rs cr)", [
            f"{b['fund'].aum_cr:,.0f}" if b["fund"].aum_cr else "-" for b in bundles
        ]),
        _row("Latest NAV", [
            f"{b['latest_nav']:.4f}" if b["latest_nav"] else "-" for b in bundles
        ]),
        _row("CAGR 1y", [
            _format_pct(b["metric"].cagr_1y) if b["metric"] else "-" for b in bundles
        ]),
        _row("CAGR 3y", [
            _format_pct(b["metric"].cagr_3y) if b["metric"] else "-" for b in bundles
        ]),
        _row("CAGR 5y", [
            _format_pct(b["metric"].cagr_5y) if b["metric"] else "-" for b in bundles
        ]),
        _row("Sharpe", [
            _format_num(b["metric"].sharpe_ratio, 2) if b["metric"] else "-" for b in bundles
        ]),
        _row("Max drawdown", [
            _format_pct(b["metric"].max_drawdown, 1) if b["metric"] else "-" for b in bundles
        ]),
    ]
    _add_heading(doc, "Side-by-side metrics", size=12)
    _make_grid_table(doc, headers, rows)

    # Normalised overlay chart
    _add_heading(doc, "Normalised performance", size=12)
    series = [(b["fund"].fund_name[:24], b["nav_history"]) for b in bundles]
    overlay_png = render_overlay_chart(series)
    _add_image_bytes(doc, overlay_png, width_cm=16.0)

    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _add_para(doc, f"Generated {ts}", size=8, color=GREY)
    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
